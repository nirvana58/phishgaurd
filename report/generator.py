"""
report/generator.py

Report generation system. Takes the structured result dict from scanner.py
and produces:
  1. A rich terminal output (always, during a scan)
  2. Saved files in any/all of: .md  .txt  .pdf  .docx

Optionally enhances the summary section with a local LLM (Ollama) when
use_llm=True was specified at scan time.

The key design principle here is a single internal `ReportData` dataclass
that all format renderers consume — add a new format later by just adding
a new `_write_<fmt>()` method without touching any other renderer.

LLM note:
  We use Ollama (local) via its REST API at http://localhost:11434.
  The user must have Ollama running with at least one model pulled.
  Default model: "llama3" — overridable via OLLAMA_MODEL env var.
  If Ollama is unreachable the report falls back gracefully (no crash).
"""

import json
import os
import textwrap
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional

import httpx
from rich.console import Console
from rich.panel import Panel
from rich.table import Table
from rich.text import Text
from rich import box

# reportlab
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table as RLTable,
    TableStyle, HRFlowable,
)

# python-docx
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPORTS_DIR = Path(__file__).resolve().parent.parent / "reports"
REPORTS_DIR.mkdir(exist_ok=True)

OLLAMA_BASE  = os.getenv("OLLAMA_HOST", "http://localhost:11434")
OLLAMA_MODEL = os.getenv("OLLAMA_MODEL", "")   # empty = auto-detect first available

console = Console()


# ── Ollama helpers ─────────────────────────────────────────────────────────────

async def _ollama_list_models() -> list[str]:
    """Return list of locally available model names from /api/tags."""
    try:
        async with httpx.AsyncClient(timeout=5.0) as client:
            resp = await client.get(f"{OLLAMA_BASE}/api/tags")
            if resp.status_code == 200:
                models = resp.json().get("models", [])
                return [m["name"] for m in models]
    except Exception:
        pass
    return []


async def _ollama_resolve_model() -> str | None:
    """
    Resolve which model to use:
      1. OLLAMA_MODEL env var if set and available
      2. First model returned by /api/tags
      3. None if Ollama is unreachable or no models installed
    """
    available = await _ollama_list_models()
    if not available:
        return None

    preferred = OLLAMA_MODEL.strip()

    # Exact match first
    if preferred and preferred in available:
        return preferred

    # Prefix match (e.g. "llama3" matches "llama3:latest" or "llama3.2:latest")
    if preferred:
        for m in available:
            if m.startswith(preferred.split(":")[0]):
                return m

    # Fall back to first available model
    return available[0]


async def _ollama_generate(model: str, prompt: str) -> str | None:
    """
    Try /api/chat first (Ollama ≥ 0.1.14), fall back to /api/generate.
    Returns the response text or None on any failure.
    """
    async with httpx.AsyncClient(timeout=60.0) as client:

        # ── Try /api/chat ──────────────────────────────────────────────────────
        try:
            resp = await client.post(
                f"{OLLAMA_BASE}/api/chat",
                json={
                    "model": model,
                    "messages": [{"role": "user", "content": prompt}],
                    "stream": False,
                },
            )
            if resp.status_code == 200:
                data = resp.json()
                return data.get("message", {}).get("content", "").strip()
            # 400 / 404 → try the older endpoint
        except httpx.TimeoutException:
            console.print("[yellow]  Ollama /api/chat timed out — trying /api/generate…[/yellow]")
        except Exception:
            pass

        # ── Fall back to /api/generate ─────────────────────────────────────────
        try:
            resp = await client.post(
                f"{OLLAMA_BASE}/api/generate",
                json={"model": model, "prompt": prompt, "stream": False},
            )
            if resp.status_code == 200:
                return resp.json().get("response", "").strip()
            body = resp.text[:300]
            console.print(f"[yellow]  Ollama /api/generate returned {resp.status_code}: {body}[/yellow]")
        except httpx.TimeoutException:
            console.print("[yellow]  Ollama /api/generate timed out.[/yellow]")
        except Exception as e:
            console.print(f"[yellow]  Ollama error: {e}[/yellow]")

    return None


# ── Internal report structure ─────────────────────────────────────────────────

@dataclass
class ReportData:
    scan_id: str
    url: str
    scanned_at: str
    verdict: str           # SAFE | SUSPICIOUS | MALICIOUS
    confidence: str        # HIGH | MEDIUM | LOW
    reasons: list[str]

    # ML section
    ml_available: bool
    ml_combined_score: float
    ml_kmeans_score: float
    ml_som_score: float
    ml_kmeans_cluster: int
    ml_models_agree: bool
    ml_version: str

    # External APIs
    vt_available: bool
    vt_malicious: int
    vt_suspicious: int
    vt_harmless: int
    vt_status: str          # "found" | "not_found" | ""

    gsb_available: bool
    gsb_is_threat: bool
    gsb_threat_types: list[str]

    # Features (for detail section)
    features: dict[str, float]

    # WHOIS section
    whois_available: bool = False
    whois_domain: str = ""
    whois_registrar: Optional[str] = None
    whois_country: Optional[str] = None
    whois_creation_date: Optional[str] = None
    whois_expiration_date: Optional[str] = None
    whois_updated_date: Optional[str] = None
    whois_age_days: Optional[int] = None
    whois_expiry_days: Optional[int] = None
    whois_ns_count: int = 0
    whois_is_new_domain: bool = False
    whois_expiring_soon: bool = False
    whois_is_ip_host: bool = False
    whois_error: Optional[str] = None

    # LLM narrative (populated later if use_llm=True)
    llm_summary: Optional[str] = None
    llm_model: Optional[str] = None


def _build_report_data(scan_id: str, result: dict) -> ReportData:
    verdict_block = result.get("verdict", {})
    ml  = result.get("ml", {})
    vt  = result.get("virustotal", {})
    gsb = result.get("safe_browsing", {})
    w   = result.get("whois", {})

    return ReportData(
        scan_id=scan_id,
        url=result.get("url", ""),
        scanned_at=result.get("scanned_at", ""),
        verdict=verdict_block.get("verdict", "UNKNOWN"),
        confidence=verdict_block.get("confidence", "UNKNOWN"),
        reasons=verdict_block.get("reasons", []),
        ml_available=ml.get("available", False),
        ml_combined_score=ml.get("combined_score", 0.0),
        ml_kmeans_score=ml.get("kmeans", {}).get("anomaly_score", 0.0),
        ml_som_score=ml.get("som", {}).get("anomaly_score", 0.0),
        ml_kmeans_cluster=ml.get("kmeans", {}).get("cluster_id", -1),
        ml_models_agree=ml.get("models_agree", False),
        ml_version=ml.get("model_version", "none"),
        vt_available=vt.get("available", False),
        vt_malicious=vt.get("malicious", 0),
        vt_suspicious=vt.get("suspicious", 0),
        vt_harmless=vt.get("harmless", 0),
        vt_status=vt.get("status", ""),
        gsb_available=gsb.get("available", False),
        gsb_is_threat=gsb.get("is_threat", False),
        gsb_threat_types=gsb.get("threat_types", []),
        features=result.get("features", {}),
        # WHOIS
        whois_available=w.get("available", False),
        whois_domain=w.get("domain", ""),
        whois_registrar=w.get("registrar"),
        whois_country=w.get("country"),
        whois_creation_date=w.get("creation_date"),
        whois_expiration_date=w.get("expiration_date"),
        whois_updated_date=w.get("updated_date"),
        whois_age_days=w.get("age_days"),
        whois_expiry_days=w.get("expiry_days"),
        whois_ns_count=w.get("ns_count", 0),
        whois_is_new_domain=w.get("is_new_domain", False),
        whois_expiring_soon=w.get("expiring_soon", False),
        whois_is_ip_host=w.get("is_ip_host", False),
        whois_error=w.get("reason") if not w.get("available") else None,
    )


# ── LLM enhancement ───────────────────────────────────────────────────────────

def _build_llm_prompt(rd: ReportData) -> str:
    """Shared prompt builder used by both single-scan and batch LLM summaries."""
    return (
        "You are a cybersecurity analyst. Analyse this URL scan result and write a "
        "concise 3-5 sentence threat summary for a security report. Be specific about "
        "what signals were found. Do not repeat numbers verbatim — interpret them. "
        "Do not use markdown headers or bullet points. Plain prose only.\n\n"
        f"URL: {rd.url}\n"
        f"Verdict: {rd.verdict} (confidence: {rd.confidence})\n"
        f"Reasons: {'; '.join(rd.reasons)}\n"
        f"ML anomaly score: {rd.ml_combined_score:.3f} "
        f"(K-means: {rd.ml_kmeans_score:.3f}, SOM: {rd.ml_som_score:.3f})\n"
        f"Models agree: {rd.ml_models_agree}\n"
        f"VirusTotal: {rd.vt_malicious} malicious, {rd.vt_suspicious} suspicious detections\n"
        f"Google Safe Browsing: "
        f"{'THREAT — ' + ', '.join(rd.gsb_threat_types) if rd.gsb_is_threat else 'no threat'}\n"
        f"Notable URL features: "
        f"entropy={rd.features.get('shannon_entropy', 0):.2f}, "
        f"typosquat_distance={rd.features.get('typosquat_distance', 0)}, "
        f"suspicious_tld={int(rd.features.get('suspicious_tld', 0))}, "
        f"has_ip_host={int(rd.features.get('has_ip_host', 0))}"
    )


async def _fetch_llm_summary(rd: ReportData) -> tuple[str | None, str | None]:
    """
    Call local Ollama to generate a human-readable threat summary.
    Returns (summary_text, model_used) or (None, None) on failure.

    Steps:
      1. Auto-detect available model (respects OLLAMA_MODEL env var)
      2. Try /api/chat, fall back to /api/generate
      3. Never raises — all errors surface as console warnings + (None, None)
    """
    model = await _ollama_resolve_model()
    if not model:
        console.print(
            "[yellow]  Ollama unavailable or no models installed.[/yellow]\n"
            "  Install Ollama: https://ollama.com\n"
            "  Pull a model:   ollama pull llama3"
        )
        return None, None

    console.print(f"[dim]  Using model: {model}[/dim]")

    text = await _ollama_generate(model, _build_llm_prompt(rd))
    return (text, model) if text else (None, None)


# ── Terminal renderer ──────────────────────────────────────────────────────────

def _verdict_color(verdict: str) -> str:
    return {"SAFE": "green", "SUSPICIOUS": "yellow", "MALICIOUS": "red"}.get(verdict, "white")


def render_terminal(rd: ReportData):
    """Print a rich formatted report to the terminal."""
    color = _verdict_color(rd.verdict)

    console.print()
    console.rule("[bold]URL Threat Scan Report[/bold]")
    console.print()

    # Header panel
    header = Text()
    header.append(f"  URL        : {rd.url}\n")
    header.append(f"  Scan ID    : {rd.scan_id}\n")
    header.append(f"  Scanned at : {rd.scanned_at}\n")
    header.append(f"  Verdict    : ", style="bold")
    header.append(f"{rd.verdict}", style=f"bold {color}")
    header.append(f"  (confidence: {rd.confidence})\n")
    console.print(Panel(header, title="[bold]Overview[/bold]", border_style=color))

    # Reasons
    console.print("\n[bold]Signal Summary[/bold]")
    for reason in rd.reasons:
        console.print(f"  [dim]•[/dim] {reason}")

    # ML section
    console.print("\n[bold]ML Anomaly Detection[/bold]")
    if rd.ml_available:
        ml_table = Table(box=box.SIMPLE, show_header=True, header_style="bold dim")
        ml_table.add_column("Model")
        ml_table.add_column("Score", justify="right")
        ml_table.add_column("Cluster/BMU", justify="right")
        ml_table.add_row(
            "K-means",
            f"{rd.ml_kmeans_score:.4f}",
            f"cluster {rd.ml_kmeans_cluster}",
        )
        ml_table.add_row(
            "SOM",
            f"{rd.ml_som_score:.4f}",
            "—",
        )
        ml_table.add_row(
            "[bold]Combined[/bold]",
            f"[bold]{rd.ml_combined_score:.4f}[/bold]",
            f"models {'[green]agree[/green]' if rd.ml_models_agree else '[yellow]disagree[/yellow]'}",
        )
        console.print(ml_table)
        console.print(f"  [dim]Model version: {rd.ml_version}[/dim]")
    else:
        console.print("  [dim]ML scoring unavailable (no trained model)[/dim]")

    # External APIs
    console.print("\n[bold]External Threat Intelligence[/bold]")
    ext_table = Table(box=box.SIMPLE, show_header=True, header_style="bold dim")
    ext_table.add_column("Source")
    ext_table.add_column("Result")
    ext_table.add_column("Detail")

    if rd.vt_available:
        vt_result = (
            f"[red]{rd.vt_malicious} malicious[/red]" if rd.vt_malicious > 0
            else "[green]clean[/green]"
        )
        ext_table.add_row(
            "VirusTotal",
            vt_result,
            f"{rd.vt_suspicious} suspicious, {rd.vt_harmless} harmless",
        )
    else:
        ext_table.add_row("VirusTotal", "[dim]unavailable[/dim]", "API key not set or timeout")

    if rd.gsb_available:
        gsb_result = (
            f"[red]THREAT[/red]" if rd.gsb_is_threat else "[green]clean[/green]"
        )
        ext_table.add_row(
            "Safe Browsing",
            gsb_result,
            ", ".join(rd.gsb_threat_types) if rd.gsb_threat_types else "—",
        )
    else:
        ext_table.add_row("Safe Browsing", "[dim]unavailable[/dim]", "API key not set or timeout")

    console.print(ext_table)

    # WHOIS Domain Intelligence
    console.print("\n[bold]WHOIS Domain Intelligence[/bold]")
    if rd.whois_available:
        whois_tbl = Table(box=box.SIMPLE, show_header=False)
        whois_tbl.add_column("Field",  style="bold dim", width=22)
        whois_tbl.add_column("Value")

        def _fmt_age(days):
            if days is None: return "—"
            if days < 30:   return f"[red bold]{days} days[/red bold] ⚑ NEWLY REGISTERED"
            if days < 90:   return f"[yellow]{days} days[/yellow] ⚑ < 90 days old"
            if days < 365:  return f"[yellow]{days} days ({days//30} months)[/yellow]"
            return f"[green]{days} days ({days//365}y {(days%365)//30}m)[/green]"

        def _fmt_expiry(days):
            if days is None: return "—"
            if days < 0:    return f"[red bold]EXPIRED {abs(days)} days ago[/red bold]"
            if days < 30:   return f"[red]{days} days remaining ⚑ EXPIRING SOON[/red]"
            if days < 180:  return f"[yellow]{days} days remaining[/yellow]"
            return f"[green]{days} days remaining[/green]"

        if rd.whois_is_ip_host:
            whois_tbl.add_row("Type", "[dim]IP address host — WHOIS not applicable[/dim]")
        else:
            whois_tbl.add_row("Domain",          rd.whois_domain or "—")
            whois_tbl.add_row("Registrar",        rd.whois_registrar or "[dim]not available[/dim]")
            whois_tbl.add_row("Country",          rd.whois_country   or "[dim]not available[/dim]")
            whois_tbl.add_row("Created",          rd.whois_creation_date or "[dim]not available[/dim]")
            whois_tbl.add_row("Expires",          rd.whois_expiration_date or "[dim]not available[/dim]")
            whois_tbl.add_row("Last Updated",     rd.whois_updated_date  or "[dim]not available[/dim]")
            whois_tbl.add_row("Domain Age",       _fmt_age(rd.whois_age_days))
            whois_tbl.add_row("Expiry",           _fmt_expiry(rd.whois_expiry_days))
            whois_tbl.add_row("Name Servers",     str(rd.whois_ns_count) if rd.whois_ns_count else "[dim]not available[/dim]")
            # Risk flags
            flags = []
            if rd.whois_is_new_domain:   flags.append("[red]⚑ Newly registered (< 30 days)[/red]")
            if rd.whois_expiring_soon:   flags.append("[yellow]⚑ Expiring soon (< 30 days)[/yellow]")
            if rd.whois_ns_count == 1:   flags.append("[yellow]⚑ Only 1 name server[/yellow]")
            if flags:
                whois_tbl.add_row("Risk Flags", "  ".join(flags))
            else:
                whois_tbl.add_row("Risk Flags", "[green]None[/green]")
    else:
        whois_tbl = Table(box=box.SIMPLE, show_header=False)
        whois_tbl.add_column("Field", width=22)
        whois_tbl.add_column("Value")
        reason = rd.whois_error or "lookup failed or not attempted"
        whois_tbl.add_row("Status", f"[dim]unavailable — {reason}[/dim]")
    console.print(whois_tbl)

    # URL Features
    console.print("\n[bold]URL Feature Breakdown[/bold]")
    feat_table = Table(box=box.SIMPLE, show_header=True, header_style="bold dim")
    feat_table.add_column("Feature")
    feat_table.add_column("Value", justify="right")
    feat_table.add_column("Flag", justify="center")

    flag_features = {
        "has_ip_host": ("IP host", "red"),
        "suspicious_tld": ("Suspicious TLD", "red"),
        "has_at_symbol": ("@ symbol", "yellow"),
    }
    for name, value in rd.features.items():
        flag = ""
        if name in flag_features and value == 1:
            label, col = flag_features[name]
            flag = f"[{col}]⚑ {label}[/{col}]"
        feat_table.add_row(name, str(round(value, 4)), flag)
    console.print(feat_table)

    # LLM summary
    if rd.llm_summary:
        console.print()
        console.print(Panel(
            rd.llm_summary,
            title=f"[bold]AI Analysis[/bold] [dim]({rd.llm_model})[/dim]",
            border_style="blue",
        ))

    console.print()
    console.rule()
    console.print()


# ── Markdown renderer ─────────────────────────────────────────────────────────

def _write_md(rd: ReportData, path: Path):
    lines = [
        f"# URL Threat Scan Report",
        f"",
        f"| Field | Value |",
        f"|-------|-------|",
        f"| **URL** | `{rd.url}` |",
        f"| **Scan ID** | `{rd.scan_id}` |",
        f"| **Scanned at** | {rd.scanned_at} |",
        f"| **Verdict** | **{rd.verdict}** |",
        f"| **Confidence** | {rd.confidence} |",
        f"",
        f"## Signal Summary",
        f"",
    ]
    for r in rd.reasons:
        lines.append(f"- {r}")

    lines += [
        f"",
        f"## ML Anomaly Detection",
        f"",
    ]
    if rd.ml_available:
        lines += [
            f"| Model | Anomaly Score |",
            f"|-------|--------------|",
            f"| K-means (cluster {rd.ml_kmeans_cluster}) | {rd.ml_kmeans_score:.4f} |",
            f"| SOM | {rd.ml_som_score:.4f} |",
            f"| **Combined** | **{rd.ml_combined_score:.4f}** |",
            f"",
            f"Models agree: {'Yes' if rd.ml_models_agree else 'No (mixed signal — lower confidence)'}  ",
            f"Model version: `{rd.ml_version}`",
        ]
    else:
        lines.append("_ML scoring unavailable (no trained model loaded)._")

    lines += [
        f"",
        f"## External Threat Intelligence",
        f"",
        f"### VirusTotal",
    ]
    if rd.vt_available:
        lines += [
            f"- Malicious detections: **{rd.vt_malicious}**",
            f"- Suspicious: {rd.vt_suspicious}",
            f"- Harmless: {rd.vt_harmless}",
            f"- Status: {rd.vt_status}",
        ]
    else:
        lines.append("_Unavailable (API key not configured or request timed out)._")

    lines += [f"", f"### Google Safe Browsing"]
    if rd.gsb_available:
        lines.append(f"- Threat detected: **{'Yes — ' + ', '.join(rd.gsb_threat_types) if rd.gsb_is_threat else 'No'}**")
    else:
        lines.append("_Unavailable (API key not configured or request timed out)._")

    # WHOIS section
    lines += [f"", f"## WHOIS Domain Intelligence", f""]
    if rd.whois_available:
        if rd.whois_is_ip_host:
            lines.append("_Host is an IP address — WHOIS domain lookup not applicable._")
        else:
            def _age_note(days):
                if days is None: return ""
                if days < 30:   return " ⚠️ **NEWLY REGISTERED**"
                if days < 90:   return " ⚠️ less than 90 days old"
                return ""
            def _exp_note(days):
                if days is None: return ""
                if days < 0:    return " ⚠️ **EXPIRED**"
                if days < 30:   return " ⚠️ **EXPIRING SOON**"
                return ""
            lines += [
                f"| Field | Value |",
                f"|-------|-------|",
                f"| Domain | `{rd.whois_domain}` |",
                f"| Registrar | {rd.whois_registrar or '—'} |",
                f"| Country | {rd.whois_country or '—'} |",
                f"| Created | {rd.whois_creation_date or '—'} |",
                f"| Expires | {rd.whois_expiration_date or '—'} |",
                f"| Last Updated | {rd.whois_updated_date or '—'} |",
                f"| Domain Age | {f'{rd.whois_age_days} days' if rd.whois_age_days is not None else '—'}"
                f"{_age_note(rd.whois_age_days)} |",
                f"| Days Until Expiry | {f'{rd.whois_expiry_days} days' if rd.whois_expiry_days is not None else '—'}"
                f"{_exp_note(rd.whois_expiry_days)} |",
                f"| Name Server Count | {rd.whois_ns_count or '—'} |",
            ]
            # Risk flags
            flags = []
            if rd.whois_is_new_domain:  flags.append("⚠️ Newly registered domain (< 30 days)")
            if rd.whois_expiring_soon:  flags.append("⚠️ Domain expiring soon (< 30 days)")
            if rd.whois_ns_count == 1:  flags.append("⚠️ Only 1 name server")
            lines += [f"", f"**WHOIS Risk Flags:**"]
            if flags:
                for flag in flags:
                    lines.append(f"- {flag}")
            else:
                lines.append("- None detected")
    else:
        reason = rd.whois_error or "lookup failed or not attempted"
        lines.append(f"_WHOIS data unavailable — {reason}_")

    lines += [f"", f"## URL Feature Breakdown", f"", f"| Feature | Value |", f"|---------|-------|"]
    for name, val in rd.features.items():
        lines.append(f"| {name} | {round(val, 4)} |")

    if rd.llm_summary:
        lines += [f"", f"## AI Analysis", f"", f"> _{rd.llm_model}_", f"", rd.llm_summary]

    path.write_text("\n".join(lines), encoding="utf-8")


# ── Plain text renderer ───────────────────────────────────────────────────────

def _write_txt(rd: ReportData, path: Path):
    sep = "=" * 60
    lines = [
        sep,
        "URL THREAT SCAN REPORT",
        sep,
        f"URL         : {rd.url}",
        f"Scan ID     : {rd.scan_id}",
        f"Scanned at  : {rd.scanned_at}",
        f"Verdict     : {rd.verdict} (confidence: {rd.confidence})",
        sep,
        "SIGNAL SUMMARY",
        sep,
    ]
    for r in rd.reasons:
        lines.append(f"  - {r}")

    lines += [sep, "ML ANOMALY DETECTION", sep]
    if rd.ml_available:
        lines += [
            f"  K-means score  : {rd.ml_kmeans_score:.4f}  (cluster {rd.ml_kmeans_cluster})",
            f"  SOM score      : {rd.ml_som_score:.4f}",
            f"  Combined score : {rd.ml_combined_score:.4f}",
            f"  Models agree   : {'Yes' if rd.ml_models_agree else 'No (mixed signal)'}",
            f"  Model version  : {rd.ml_version}",
        ]
    else:
        lines.append("  ML scoring unavailable.")

    lines += [sep, "EXTERNAL THREAT INTELLIGENCE", sep]
    if rd.vt_available:
        lines += [
            f"  VirusTotal  — malicious: {rd.vt_malicious}, suspicious: {rd.vt_suspicious}, harmless: {rd.vt_harmless}",
        ]
    else:
        lines.append("  VirusTotal  — unavailable")

    if rd.gsb_available:
        threat_str = "THREAT: " + ", ".join(rd.gsb_threat_types) if rd.gsb_is_threat else "clean"
        lines.append(f"  Safe Browsing — {threat_str}")
    else:
        lines.append("  Safe Browsing — unavailable")

    lines += [sep, "WHOIS DOMAIN INTELLIGENCE", sep]
    if rd.whois_available:
        if rd.whois_is_ip_host:
            lines.append("  IP address host — WHOIS not applicable.")
        else:
            lines += [
                f"  Domain       : {rd.whois_domain or '—'}",
                f"  Registrar    : {rd.whois_registrar or '—'}",
                f"  Country      : {rd.whois_country or '—'}",
                f"  Created      : {rd.whois_creation_date or '—'}",
                f"  Expires      : {rd.whois_expiration_date or '—'}",
                f"  Last Updated : {rd.whois_updated_date or '—'}",
                f"  Domain Age   : {f'{rd.whois_age_days} days' if rd.whois_age_days is not None else '—'}"
                + (" [!] NEWLY REGISTERED" if rd.whois_is_new_domain else ""),
                f"  Until Expiry : {f'{rd.whois_expiry_days} days' if rd.whois_expiry_days is not None else '—'}"
                + (" [!] EXPIRING SOON" if rd.whois_expiring_soon else ""),
                f"  Name Servers : {rd.whois_ns_count or '—'}",
            ]
            flags = []
            if rd.whois_is_new_domain:  flags.append("[!] Newly registered domain (< 30 days)")
            if rd.whois_expiring_soon:  flags.append("[!] Domain expiring soon (< 30 days)")
            if rd.whois_ns_count == 1:  flags.append("[!] Only 1 name server")
            lines.append(f"  Risk Flags   : {', '.join(flags) if flags else 'None'}")
    else:
        lines.append(f"  Status       : unavailable — {rd.whois_error or 'lookup failed'}")

    lines += [sep, "URL FEATURES", sep]
    for name, val in rd.features.items():
        lines.append(f"  {name:<26}: {round(val, 4)}")

    if rd.llm_summary:
        lines += [sep, f"AI ANALYSIS ({rd.llm_model})", sep]
        wrapped = textwrap.fill(rd.llm_summary, width=60)
        lines.append(wrapped)

    lines.append(sep)
    path.write_text("\n".join(lines), encoding="utf-8")


# ── PDF renderer ──────────────────────────────────────────────────────────────

def _write_pdf(rd: ReportData, path: Path):
    doc = SimpleDocTemplate(
        str(path), pagesize=A4,
        leftMargin=20*mm, rightMargin=20*mm,
        topMargin=20*mm, bottomMargin=20*mm,
    )
    styles = getSampleStyleSheet()
    VERDICT_COLOR = {
        "SAFE": colors.HexColor("#1a7f37"),
        "SUSPICIOUS": colors.HexColor("#bf8700"),
        "MALICIOUS": colors.HexColor("#cf222e"),
    }
    v_color = VERDICT_COLOR.get(rd.verdict, colors.black)

    title_style = ParagraphStyle("ReportTitle", parent=styles["Title"], fontSize=18, spaceAfter=6)
    h1_style    = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=13, spaceBefore=14, spaceAfter=4)
    normal      = styles["Normal"]
    small       = ParagraphStyle("Small", parent=normal, fontSize=8, textColor=colors.grey)
    verdict_style = ParagraphStyle(
        "Verdict", parent=styles["Normal"], fontSize=20, textColor=v_color,
        fontName="Helvetica-Bold", spaceAfter=4,
    )

    story = []

    story.append(Paragraph("URL Threat Scan Report", title_style))
    story.append(Paragraph(f"Scan ID: {rd.scan_id}", small))
    story.append(Paragraph(f"Scanned: {rd.scanned_at}", small))
    story.append(Spacer(1, 6*mm))
    story.append(Paragraph(f"{rd.verdict}", verdict_style))
    story.append(Paragraph(f"Confidence: {rd.confidence}", normal))
    story.append(Paragraph(f"URL: <font name='Courier'>{rd.url}</font>", normal))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.lightgrey, spaceAfter=4))

    story.append(Paragraph("Signal Summary", h1_style))
    for r in rd.reasons:
        story.append(Paragraph(f"• {r}", normal))
    story.append(Spacer(1, 4*mm))

    story.append(Paragraph("ML Anomaly Detection", h1_style))
    if rd.ml_available:
        ml_data = [
            ["Model", "Score", "Detail"],
            ["K-means", f"{rd.ml_kmeans_score:.4f}", f"cluster {rd.ml_kmeans_cluster}"],
            ["SOM", f"{rd.ml_som_score:.4f}", "—"],
            ["Combined", f"{rd.ml_combined_score:.4f}", "agree" if rd.ml_models_agree else "disagree"],
        ]
        t = RLTable(ml_data, colWidths=[60*mm, 40*mm, 60*mm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#fafafa")]),
        ]))
        story.append(t)
        story.append(Paragraph(f"Model version: {rd.ml_version}", small))
    else:
        story.append(Paragraph("ML scoring unavailable (no trained model).", normal))
    story.append(Spacer(1, 4*mm))

    story.append(Paragraph("External Threat Intelligence", h1_style))
    ext_data = [["Source", "Result", "Detail"]]
    if rd.vt_available:
        ext_data.append(["VirusTotal", f"{rd.vt_malicious} malicious", f"{rd.vt_suspicious} susp / {rd.vt_harmless} harmless"])
    else:
        ext_data.append(["VirusTotal", "unavailable", ""])
    if rd.gsb_available:
        ext_data.append(["Safe Browsing", "THREAT" if rd.gsb_is_threat else "clean", ", ".join(rd.gsb_threat_types)])
    else:
        ext_data.append(["Safe Browsing", "unavailable", ""])
    t2 = RLTable(ext_data, colWidths=[50*mm, 50*mm, 60*mm])
    t2.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#fafafa")]),
    ]))
    story.append(t2)
    story.append(Spacer(1, 4*mm))

    # WHOIS section
    story.append(Paragraph("WHOIS Domain Intelligence", h1_style))
    RISK_COLOR = colors.HexColor("#cf222e")
    WARN_COLOR = colors.HexColor("#bf8700")
    OK_COLOR   = colors.HexColor("#1a7f37")

    if rd.whois_available:
        if rd.whois_is_ip_host:
            story.append(Paragraph("Host is an IP address — WHOIS domain lookup not applicable.", normal))
        else:
            def _age_str(days):
                if days is None: return "—"
                return f"{days} days"
            def _exp_str(days):
                if days is None: return "—"
                if days < 0:    return f"EXPIRED ({abs(days)} days ago)"
                return f"{days} days"

            whois_data = [
                ["Field", "Value", "Risk"],
                ["Domain",          rd.whois_domain or "—",               ""],
                ["Registrar",       rd.whois_registrar or "—",            ""],
                ["Country",         rd.whois_country or "—",              ""],
                ["Created",         rd.whois_creation_date or "—",        "⚠ NEW" if rd.whois_is_new_domain else ""],
                ["Expires",         rd.whois_expiration_date or "—",      "⚠ SOON" if rd.whois_expiring_soon else ""],
                ["Last Updated",    rd.whois_updated_date or "—",         ""],
                ["Domain Age",      _age_str(rd.whois_age_days),          "⚠ < 30 days" if rd.whois_is_new_domain else ("⚠ < 90 days" if rd.whois_age_days and rd.whois_age_days < 90 else "")],
                ["Until Expiry",    _exp_str(rd.whois_expiry_days),       "⚠ EXPIRING" if rd.whois_expiring_soon else ""],
                ["Name Servers",    str(rd.whois_ns_count) if rd.whois_ns_count else "—",
                                    "⚠ Only 1 NS" if rd.whois_ns_count == 1 else ""],
            ]
            tw = RLTable(whois_data, colWidths=[38*mm, 90*mm, 32*mm])
            tw.setStyle(TableStyle([
                ("BACKGROUND",    (0,0),(-1,0),  colors.HexColor("#f0f0f0")),
                ("FONTNAME",      (0,0),(-1,0),  "Helvetica-Bold"),
                ("FONTNAME",      (0,1),(0,-1),  "Helvetica-Bold"),
                ("FONTSIZE",      (0,0),(-1,-1), 8),
                ("GRID",          (0,0),(-1,-1), 0.25, colors.lightgrey),
                ("ROWBACKGROUNDS",(0,1),(-1,-1), [colors.white, colors.HexColor("#fafafa")]),
                ("TEXTCOLOR",     (2,1),(-1,-1), WARN_COLOR),
                ("FONTNAME",      (2,1),(-1,-1), "Helvetica-Bold"),
                ("VALIGN",        (0,0),(-1,-1), "MIDDLE"),
                ("TOPPADDING",    (0,0),(-1,-1), 5),
                ("BOTTOMPADDING", (0,0),(-1,-1), 5),
            ]))
            story.append(tw)

            # Risk flags summary
            flags = []
            if rd.whois_is_new_domain:  flags.append("Newly registered domain (< 30 days)")
            if rd.whois_expiring_soon:  flags.append("Domain expiring soon (< 30 days)")
            if rd.whois_ns_count == 1:  flags.append("Only 1 name server")
            if flags:
                story.append(Spacer(1, 2*mm))
                story.append(Paragraph(
                    "<b>WHOIS Risk Flags:</b> " + " | ".join(f"⚠ {f}" for f in flags),
                    ParagraphStyle("wf", parent=normal, fontSize=8, textColor=WARN_COLOR)
                ))
            else:
                story.append(Spacer(1, 2*mm))
                story.append(Paragraph(
                    "WHOIS Risk Flags: None detected",
                    ParagraphStyle("wf_ok", parent=normal, fontSize=8, textColor=OK_COLOR)
                ))
    else:
        story.append(Paragraph(
            f"WHOIS data unavailable — {rd.whois_error or 'lookup failed or not attempted'}",
            ParagraphStyle("wna", parent=normal, fontSize=9, textColor=colors.grey)
        ))
    story.append(Spacer(1, 4*mm))

    story.append(Paragraph("URL Feature Breakdown", h1_style))
    feat_data = [["Feature", "Value"]] + [[k, str(round(v, 4))] for k, v in rd.features.items()]
    t3 = RLTable(feat_data, colWidths=[90*mm, 70*mm])
    t3.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#fafafa")]),
    ]))
    story.append(t3)

    if rd.llm_summary:
        story.append(Spacer(1, 4*mm))
        story.append(Paragraph(f"AI Analysis ({rd.llm_model})", h1_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#0969da"), spaceAfter=4))
        story.append(Paragraph(rd.llm_summary, normal))

    doc.build(story)


# ── DOCX renderer ─────────────────────────────────────────────────────────────

def _write_docx(rd: ReportData, path: Path):
    doc = Document()

    # Title
    title = doc.add_heading("URL Threat Scan Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph()
    meta.add_run(f"Scan ID: ").bold = False
    meta.add_run(rd.scan_id).font.size = Pt(9)
    p2 = doc.add_paragraph()
    p2.add_run(f"Scanned: {rd.scanned_at}").font.size = Pt(9)

    # Verdict
    doc.add_heading("Verdict", level=1)
    VERDICT_RGB = {
        "SAFE": RGBColor(0x1a, 0x7f, 0x37),
        "SUSPICIOUS": RGBColor(0xbf, 0x87, 0x00),
        "MALICIOUS": RGBColor(0xcf, 0x22, 0x2e),
    }
    vp = doc.add_paragraph()
    run = vp.add_run(rd.verdict)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = VERDICT_RGB.get(rd.verdict, RGBColor(0, 0, 0))
    doc.add_paragraph(f"Confidence: {rd.confidence}")
    url_p = doc.add_paragraph()
    url_run = url_p.add_run(f"URL: {rd.url}")
    url_run.font.name = "Courier New"
    url_run.font.size = Pt(9)

    # Reasons
    doc.add_heading("Signal Summary", level=1)
    for r in rd.reasons:
        doc.add_paragraph(r, style="List Bullet")

    # ML
    doc.add_heading("ML Anomaly Detection", level=1)
    if rd.ml_available:
        ml_tbl = doc.add_table(rows=4, cols=3)
        ml_tbl.style = "Table Grid"
        headers = ["Model", "Score", "Detail"]
        for i, h in enumerate(headers):
            cell = ml_tbl.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
        rows = [
            ["K-means", f"{rd.ml_kmeans_score:.4f}", f"cluster {rd.ml_kmeans_cluster}"],
            ["SOM",     f"{rd.ml_som_score:.4f}",    "—"],
            ["Combined",f"{rd.ml_combined_score:.4f}", "agree" if rd.ml_models_agree else "disagree (mixed signal)"],
        ]
        for ri, row_data in enumerate(rows, start=1):
            for ci, val in enumerate(row_data):
                ml_tbl.rows[ri].cells[ci].text = val
        doc.add_paragraph(f"Model version: {rd.ml_version}").runs[0].font.size = Pt(8)
    else:
        doc.add_paragraph("ML scoring unavailable (no trained model loaded).")

    # External APIs
    doc.add_heading("External Threat Intelligence", level=1)
    ext_tbl = doc.add_table(rows=3, cols=3)
    ext_tbl.style = "Table Grid"
    for i, h in enumerate(["Source", "Result", "Detail"]):
        ext_tbl.rows[0].cells[i].text = h
        ext_tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    if rd.vt_available:
        ext_tbl.rows[1].cells[0].text = "VirusTotal"
        ext_tbl.rows[1].cells[1].text = f"{rd.vt_malicious} malicious"
        ext_tbl.rows[1].cells[2].text = f"{rd.vt_suspicious} suspicious, {rd.vt_harmless} harmless"
    else:
        ext_tbl.rows[1].cells[0].text = "VirusTotal"
        ext_tbl.rows[1].cells[1].text = "unavailable"

    if rd.gsb_available:
        ext_tbl.rows[2].cells[0].text = "Google Safe Browsing"
        ext_tbl.rows[2].cells[1].text = "THREAT" if rd.gsb_is_threat else "clean"
        ext_tbl.rows[2].cells[2].text = ", ".join(rd.gsb_threat_types) if rd.gsb_threat_types else "—"
    else:
        ext_tbl.rows[2].cells[0].text = "Google Safe Browsing"
        ext_tbl.rows[2].cells[1].text = "unavailable"

    # WHOIS section
    doc.add_heading("WHOIS Domain Intelligence", level=1)
    if rd.whois_available:
        if rd.whois_is_ip_host:
            doc.add_paragraph("Host is an IP address — WHOIS domain lookup not applicable.")
        else:
            whois_rows = [
                ("Domain",         rd.whois_domain or "—",                False),
                ("Registrar",      rd.whois_registrar or "—",             False),
                ("Country",        rd.whois_country or "—",               False),
                ("Created",        rd.whois_creation_date or "—",         False),
                ("Expires",        rd.whois_expiration_date or "—",       rd.whois_expiring_soon),
                ("Last Updated",   rd.whois_updated_date or "—",          False),
                ("Domain Age",     f"{rd.whois_age_days} days" if rd.whois_age_days is not None else "—",
                                   rd.whois_is_new_domain),
                ("Until Expiry",   f"{rd.whois_expiry_days} days" if rd.whois_expiry_days is not None else "—",
                                   rd.whois_expiring_soon),
                ("Name Servers",   str(rd.whois_ns_count) if rd.whois_ns_count else "—",
                                   rd.whois_ns_count == 1),
            ]
            wtbl = doc.add_table(rows=len(whois_rows) + 1, cols=3)
            wtbl.style = "Table Grid"
            for ci, h in enumerate(["Field", "Value", "Risk Flag"]):
                wtbl.rows[0].cells[ci].text = h
                wtbl.rows[0].cells[ci].paragraphs[0].runs[0].bold = True
            WARN = RGBColor(0xbf, 0x87, 0x00)
            for ri, (field, value, is_risk) in enumerate(whois_rows, 1):
                wtbl.rows[ri].cells[0].text = field
                wtbl.rows[ri].cells[1].text = value
                flag_text = "⚠ Risk detected" if is_risk else ""
                wtbl.rows[ri].cells[2].text = flag_text
                if is_risk and wtbl.rows[ri].cells[2].paragraphs[0].runs:
                    wtbl.rows[ri].cells[2].paragraphs[0].runs[0].font.color.rgb = WARN
                    wtbl.rows[ri].cells[2].paragraphs[0].runs[0].bold = True

            # Risk summary paragraph
            flags = []
            if rd.whois_is_new_domain:  flags.append("Newly registered (< 30 days)")
            if rd.whois_expiring_soon:  flags.append("Expiring soon (< 30 days)")
            if rd.whois_ns_count == 1:  flags.append("Only 1 name server")
            p = doc.add_paragraph()
            r = p.add_run("WHOIS Risk Flags: " + (", ".join(flags) if flags else "None detected"))
            r.font.size = Pt(9)
            if flags:
                r.font.color.rgb = WARN
                r.bold = True
    else:
        doc.add_paragraph(
            f"WHOIS data unavailable — {rd.whois_error or 'lookup failed or not attempted'}"
        )

    # Features table
    doc.add_heading("URL Feature Breakdown", level=1)
    feat_tbl = doc.add_table(rows=len(rd.features) + 1, cols=2)
    feat_tbl.style = "Table Grid"
    feat_tbl.rows[0].cells[0].text = "Feature"
    feat_tbl.rows[0].cells[1].text = "Value"
    feat_tbl.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    feat_tbl.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    for ri, (k, v) in enumerate(rd.features.items(), start=1):
        feat_tbl.rows[ri].cells[0].text = k
        feat_tbl.rows[ri].cells[1].text = str(round(v, 4))

    # LLM section
    if rd.llm_summary:
        doc.add_heading(f"AI Analysis ({rd.llm_model})", level=1)
        doc.add_paragraph(rd.llm_summary)

    doc.save(str(path))


# ── Public API ────────────────────────────────────────────────────────────────

async def generate_report(
    scan_id: str,
    result: dict,
    formats: list[str] = None,    # ["md", "txt", "pdf", "docx"] or subset
    show_terminal: bool = True,
) -> dict[str, Path]:
    """
    Main entry point called by the CLI after a scan completes.

    Args:
        scan_id:       UUID of the completed scan.
        result:        The result dict from scan_jobs.result_json.
        formats:       List of file formats to save. None = all four.
        show_terminal: Whether to print the rich terminal report.

    Returns:
        dict mapping format -> saved Path (only for formats that were saved).
    """
    if formats is None:
        formats = ["md", "txt", "pdf", "docx"]

    rd = _build_report_data(scan_id, result)

    # LLM enhancement
    if result.get("use_llm"):
        console.print("[dim]Generating AI summary via Ollama…[/dim]")
        rd.llm_summary, rd.llm_model = await _fetch_llm_summary(rd)
        if not rd.llm_summary:
            console.print("[yellow]  Skipping AI summary — see warnings above.[/yellow]")

    # Terminal
    if show_terminal:
        render_terminal(rd)

    # File outputs
    saved: dict[str, Path] = {}
    writers = {
        "md":   _write_md,
        "txt":  _write_txt,
        "pdf":  _write_pdf,
        "docx": _write_docx,
    }

    for fmt in formats:
        if fmt not in writers:
            console.print(f"[yellow]Unknown format '{fmt}' — skipping.[/yellow]")
            continue
        out_path = REPORTS_DIR / f"{scan_id}.{fmt}"
        try:
            writers[fmt](rd, out_path)
            saved[fmt] = out_path
            console.print(f"[dim]  Saved {fmt.upper()} report → {out_path}[/dim]")
        except Exception as e:
            console.print(f"[red]  Failed to write {fmt} report: {e}[/red]")

    return saved


# ── Batch report ───────────────────────────────────────────────────────────────

async def generate_batch_report(
    batch_id: str,
    results: list[dict],
    formats: list[str] = None,
    show_terminal: bool = True,
) -> dict[str, Path]:
    """
    Generate a single consolidated report for a batch scan.
    One file per format, named batch_<batch_id>.<fmt>.
    Terminal output shows a summary table + per-URL verdict breakdown.
    LLM summary is generated per-URL if use_llm=True is present in any result.
    """
    if formats is None:
        formats = ["md", "txt", "pdf", "docx"]

    if not results:
        console.print("[yellow]Batch report: no results to report.[/yellow]")
        return {}

    scanned_at = results[0].get("scanned_at", "")
    total      = len(results)
    counts     = {"SAFE": 0, "SUSPICIOUS": 0, "MALICIOUS": 0, "UNKNOWN": 0}
    use_llm    = any(r.get("use_llm") for r in results)
    rows       = []

    # ── LLM summaries (one per URL if enabled) ────────────────────────────────
    llm_summaries: dict[str, str] = {}
    llm_model_name: str | None = None

    if use_llm:
        console.print(f"[dim]Generating LLM summaries for {total} URLs via Ollama…[/dim]")
        model = await _ollama_resolve_model()
        if not model:
            console.print("[yellow]  Ollama unavailable — skipping LLM summaries.[/yellow]")
        else:
            llm_model_name = model
            for i, r in enumerate(results, 1):
                rd = _build_report_data(f"batch-{i}", r)
                console.print(f"  [dim]({i}/{total}) {r.get('url','')[:60]}[/dim]")
                text = await _ollama_generate(model, _build_llm_prompt(rd))
                if text:
                    llm_summaries[r.get("url", "")] = text

    for r in results:
        v_block    = r.get("verdict", {})
        verdict    = v_block.get("verdict", "UNKNOWN")
        confidence = v_block.get("confidence", "—")
        reasons    = v_block.get("reasons", [])
        url        = r.get("url", "")
        counts[verdict] = counts.get(verdict, 0) + 1
        rows.append({
            "url": url, "verdict": verdict,
            "confidence": confidence, "reasons": reasons,
            "ml_score": r.get("ml", {}).get("combined_score", 0.0),
            "vt_malicious": r.get("virustotal", {}).get("malicious", 0),
            "gsb_threat": r.get("safe_browsing", {}).get("is_threat", False),
            "llm_summary": llm_summaries.get(url),
        })

    # ── Terminal ──────────────────────────────────────────────────────────────
    if show_terminal:
        from rich.table import Table as RichTable
        from rich import box as rbox

        console.print()
        console.rule(f"[bold]Batch Scan Report — {batch_id[:8]}…[/bold]")
        console.print(f"\n  Total URLs   : {total}")
        console.print(f"  [green]Safe         : {counts.get('SAFE',0)}[/green]")
        console.print(f"  [yellow]Suspicious   : {counts.get('SUSPICIOUS',0)}[/yellow]")
        console.print(f"  [red]Malicious    : {counts.get('MALICIOUS',0)}[/red]")
        console.print()

        tbl = RichTable(box=rbox.ROUNDED, show_lines=True, title="Per-URL Results")
        tbl.add_column("#",          width=4,  justify="right")
        tbl.add_column("URL",        max_width=50, no_wrap=True)
        tbl.add_column("Verdict",    width=12)
        tbl.add_column("Confidence", width=10)
        tbl.add_column("ML Score",   width=9,  justify="right")
        tbl.add_column("VT Mal",     width=7,  justify="right")
        tbl.add_column("GSB",        width=6,  justify="center")

        COLOR = {"SAFE": "green", "SUSPICIOUS": "yellow", "MALICIOUS": "red"}
        for i, row in enumerate(rows, 1):
            col = COLOR.get(row["verdict"], "white")
            tbl.add_row(
                str(i), row["url"],
                f"[{col}]{row['verdict']}[/{col}]",
                row["confidence"],
                f"{row['ml_score']:.3f}",
                str(row["vt_malicious"]),
                "[red]✗[/red]" if row["gsb_threat"] else "[green]✓[/green]",
            )
        console.print(tbl)
        console.print()

    # ── File outputs ──────────────────────────────────────────────────────────
    saved: dict[str, Path] = {}
    prefix = f"batch_{batch_id}"

    # Markdown
    if "md" in formats:
        md_lines = [
            f"# Batch Scan Report",
            f"",
            f"| Field | Value |",
            f"|-------|-------|",
            f"| **Batch ID** | `{batch_id}` |",
            f"| **Scanned at** | {scanned_at} |",
            f"| **Total URLs** | {total} |",
            f"| **Safe** | {counts.get('SAFE',0)} |",
            f"| **Suspicious** | {counts.get('SUSPICIOUS',0)} |",
            f"| **Malicious** | {counts.get('MALICIOUS',0)} |",
            f"",
            f"## Per-URL Results",
            f"",
            f"| # | URL | Verdict | Confidence | ML Score | VT Malicious | GSB Threat |",
            f"|---|-----|---------|------------|----------|-------------|-----------|",
        ]
        for i, row in enumerate(rows, 1):
            gsb = "Yes" if row["gsb_threat"] else "No"
            md_lines.append(
                f"| {i} | `{row['url']}` | **{row['verdict']}** | "
                f"{row['confidence']} | {row['ml_score']:.3f} | "
                f"{row['vt_malicious']} | {gsb} |"
            )
        md_lines += ["", "## Detail per URL", ""]
        for i, (row, result) in enumerate(zip(rows, results), 1):
            md_lines += [
                f"### {i}. `{row['url']}`",
                f"- **Verdict**: {row['verdict']} ({row['confidence']})",
            ]
            for reason in row["reasons"]:
                md_lines.append(f"- {reason}")
            if row.get("llm_summary"):
                md_lines += [
                    "",
                    f"> **AI Analysis** _{llm_model_name}_",
                    f"> {row['llm_summary']}",
                ]
            md_lines.append("")

        p = REPORTS_DIR / f"{prefix}.md"
        p.write_text("\n".join(md_lines), encoding="utf-8")
        saved["md"] = p
        console.print(f"[dim]  Saved MD  → {p}[/dim]")

    # Plain text
    if "txt" in formats:
        sep = "=" * 60
        txt_lines = [
            sep, "BATCH SCAN REPORT", sep,
            f"Batch ID    : {batch_id}",
            f"Scanned at  : {scanned_at}",
            f"Total URLs  : {total}",
            f"Safe        : {counts.get('SAFE',0)}",
            f"Suspicious  : {counts.get('SUSPICIOUS',0)}",
            f"Malicious   : {counts.get('MALICIOUS',0)}",
            sep, "PER-URL RESULTS", sep,
        ]
        for i, row in enumerate(rows, 1):
            txt_lines += [
                f"  {i:3}. {row['url']}",
                f"       Verdict    : {row['verdict']} ({row['confidence']})",
                f"       ML Score   : {row['ml_score']:.4f}",
                f"       VT Mal     : {row['vt_malicious']}",
                f"       GSB Threat : {'Yes' if row['gsb_threat'] else 'No'}",
                "",
            ]
        p = REPORTS_DIR / f"{prefix}.txt"
        p.write_text("\n".join(txt_lines), encoding="utf-8")
        saved["txt"] = p
        console.print(f"[dim]  Saved TXT → {p}[/dim]")

    # PDF
    if "pdf" in formats:
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm
            from reportlab.lib import colors
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, Table as RLTable,
                TableStyle, HRFlowable,
            )
            pdf_path = REPORTS_DIR / f"{prefix}.pdf"
            doc = SimpleDocTemplate(str(pdf_path), pagesize=A4,
                                    leftMargin=18*mm, rightMargin=18*mm,
                                    topMargin=18*mm, bottomMargin=18*mm)
            styles  = getSampleStyleSheet()
            normal  = styles["Normal"]
            title_s = ParagraphStyle("T", parent=styles["Title"], fontSize=16, spaceAfter=6)
            h1_s    = ParagraphStyle("H", parent=styles["Heading1"], fontSize=12, spaceBefore=10, spaceAfter=4)
            small   = ParagraphStyle("S", parent=normal, fontSize=8, textColor=colors.grey)
            C = {"SAFE": colors.HexColor("#1a7f37"),
                 "SUSPICIOUS": colors.HexColor("#bf8700"),
                 "MALICIOUS": colors.HexColor("#cf222e")}

            story = [
                Paragraph("Batch Scan Report", title_s),
                Paragraph(f"Batch ID: {batch_id}", small),
                Paragraph(f"Scanned: {scanned_at}  |  Total: {total}  |  "
                          f"Safe: {counts.get('SAFE',0)}  Suspicious: {counts.get('SUSPICIOUS',0)}  "
                          f"Malicious: {counts.get('MALICIOUS',0)}", normal),
                Spacer(1, 6*mm),
                Paragraph("Per-URL Results", h1_s),
            ]
            tbl_data = [["#", "URL", "Verdict", "Conf.", "ML", "VT", "GSB"]]
            for i, row in enumerate(rows, 1):
                tbl_data.append([
                    str(i),
                    Paragraph(f'<font size="7" name="Courier">{row["url"][:60]}</font>', normal),
                    Paragraph(f'<font color="{C.get(row["verdict"], colors.black).hexval()}">'
                              f'<b>{row["verdict"]}</b></font>', normal),
                    row["confidence"],
                    f'{row["ml_score"]:.3f}',
                    str(row["vt_malicious"]),
                    "Y" if row["gsb_threat"] else "N",
                ])
            t = RLTable(tbl_data, colWidths=[10*mm, 75*mm, 28*mm, 18*mm, 16*mm, 10*mm, 10*mm])
            t.setStyle(TableStyle([
                ("BACKGROUND",   (0,0),(-1,0),  colors.HexColor("#f0f0f0")),
                ("FONTNAME",     (0,0),(-1,0),  "Helvetica-Bold"),
                ("FONTSIZE",     (0,0),(-1,-1), 8),
                ("GRID",         (0,0),(-1,-1), 0.25, colors.lightgrey),
                ("ROWBACKGROUNDS",(0,1),(-1,-1),[colors.white, colors.HexColor("#fafafa")]),
                ("VALIGN",       (0,0),(-1,-1), "MIDDLE"),
            ]))
            story.append(t)
            doc.build(story)
            saved["pdf"] = pdf_path
            console.print(f"[dim]  Saved PDF → {pdf_path}[/dim]")
        except Exception as e:
            console.print(f"[red]  PDF batch report failed: {e}[/red]")

    # DOCX
    if "docx" in formats:
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            doc = Document()
            doc.add_heading("Batch Scan Report", 0)
            doc.add_paragraph(f"Batch ID: {batch_id}")
            doc.add_paragraph(f"Scanned: {scanned_at}")
            doc.add_paragraph(f"Total: {total}  |  Safe: {counts.get('SAFE',0)}  "
                               f"Suspicious: {counts.get('SUSPICIOUS',0)}  "
                               f"Malicious: {counts.get('MALICIOUS',0)}")
            doc.add_heading("Per-URL Results", 1)
            tbl = doc.add_table(rows=len(rows)+1, cols=6)
            tbl.style = "Table Grid"
            for i, h in enumerate(["URL","Verdict","Confidence","ML Score","VT Mal","GSB"]):
                tbl.rows[0].cells[i].text = h
                tbl.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            VRGB = {"SAFE": RGBColor(0x1a,0x7f,0x37),
                    "SUSPICIOUS": RGBColor(0xbf,0x87,0x00),
                    "MALICIOUS": RGBColor(0xcf,0x22,0x2e)}
            for ri, row in enumerate(rows, 1):
                tbl.rows[ri].cells[0].text = row["url"]
                c = tbl.rows[ri].cells[1]
                c.text = row["verdict"]
                if c.paragraphs[0].runs:
                    c.paragraphs[0].runs[0].font.color.rgb = VRGB.get(row["verdict"], RGBColor(0,0,0))
                tbl.rows[ri].cells[2].text = row["confidence"]
                tbl.rows[ri].cells[3].text = f"{row['ml_score']:.4f}"
                tbl.rows[ri].cells[4].text = str(row["vt_malicious"])
                tbl.rows[ri].cells[5].text = "Yes" if row["gsb_threat"] else "No"
            docx_path = REPORTS_DIR / f"{prefix}.docx"
            doc.save(str(docx_path))
            saved["docx"] = docx_path
            console.print(f"[dim]  Saved DOCX → {docx_path}[/dim]")
        except Exception as e:
            console.print(f"[red]  DOCX batch report failed: {e}[/red]")

    return saved
